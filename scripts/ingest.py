"""
scripts/ingest.py — Read .docx or .txt → chunk → embed → store in Supabase pgvector

Usage:
    python -m scripts.ingest
    python -m scripts.ingest --path /custom/path/to/file.txt
"""
from __future__ import annotations

import argparse
import logging
import sys
import time
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

logging.basicConfig(
    level   = logging.INFO,
    format  = "%(asctime)s | %(levelname)s | %(message)s",
    datefmt = "%H:%M:%S",
)
logger = logging.getLogger(__name__)


def read_source_file(path: Path) -> str:
    """Read .docx or .txt knowledge base file."""
    if path.suffix.lower() == '.txt':
        text = path.read_text(encoding='utf-8')
        logger.info(f"Extracted {len(text):,} characters from {path.name}")
        return text
    else:
        from docx import Document
        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(' | '.join(cells))
        raw = '\n\n'.join(parts)
        logger.info(f"Extracted {len(raw):,} characters from {path.name}")
        return raw


def chunk_text(text: str, chunk_size: int, overlap: int) -> list[dict]:
    """Split text into overlapping chunks with section metadata."""
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    splitter = RecursiveCharacterTextSplitter(
        chunk_size    = chunk_size,
        chunk_overlap = overlap,
        separators    = ["\n\n", "\n", ". ", " ", ""],
    )
    raw_chunks = splitter.split_text(text)

    section_markers = [
        "COMPANY", "PRODUCT", "PRICING", "DELIVERY", "PAYMENT",
        "QUALITY", "HOW TO", "CONTACT", "FAQ", "SIGNAGE",
        "BRAND", "EVENT", "FLYER", "POSTER", "BANNER", "STAMP",
        "BROCHURE", "ENVELOPE", "FOLDER", "MERCHANDISE", "OVERVIEW",
        "SECTION",
    ]

    result: list[dict] = []
    current_section    = "General"

    for i, chunk in enumerate(raw_chunks):
        for line in chunk.split("\n"):
            line = line.strip()
            if line and len(line) < 80:
                if any(line.upper().startswith(kw) for kw in section_markers):
                    current_section = line[:60]
                    break
        result.append({
            "content":      chunk,
            "section":      current_section,
            "chunk_index":  i,
            "total_chunks": len(raw_chunks),
        })

    logger.info(f"Created {len(result)} chunks")
    return result


def run_ingestion(docx_path: Path) -> int:
    """
    Full pipeline:
      1. Read file
      2. Chunk
      3. Embed with Gemini API
      4. Clear old Supabase data
      5. Insert new chunks + embeddings
    Returns total chunks stored.
    """
    from app.embeddings import embed_texts
    from app.database   import (
        clear_documents,
        insert_document_chunks,
        document_count,
    )

    if not docx_path.exists():
        raise FileNotFoundError(
            f"\nFile not found: {docx_path.resolve()}\n"
            f"   Place your knowledge base file in the knowledge_base/ folder.\n"
        )

    from app.config import get_settings
    settings = get_settings()

    # 1. Read
    raw_text = read_source_file(docx_path)

    # 2. Chunk
    chunks = chunk_text(raw_text, settings.chunk_size, settings.chunk_overlap)
    texts  = [c["content"] for c in chunks]

    # 3. Embed
    logger.info(f"Embedding {len(texts)} chunks with Gemini API...")
    embeddings = embed_texts(texts)
    logger.info("Embeddings complete")

    # 4. Clear old data
    logger.info("Clearing existing Supabase documents...")
    clear_documents()

    # 5. Insert in batches of 50
    BATCH          = 50
    total_inserted = 0

    for i in range(0, len(chunks), BATCH):
        batch_chunks     = chunks[i : i + BATCH]
        batch_embeddings = embeddings[i : i + BATCH]

        rows = [
            {
                "content":   batch_chunks[j]["content"],
                "metadata":  {
                    "section":      batch_chunks[j]["section"],
                    "chunk_index":  batch_chunks[j]["chunk_index"],
                    "total_chunks": batch_chunks[j]["total_chunks"],
                    "source":       docx_path.name,
                },
                "embedding": batch_embeddings[j],
            }
            for j in range(len(batch_chunks))
        ]

        insert_document_chunks(rows)
        total_inserted += len(rows)
        logger.info(f"  Stored {total_inserted}/{len(chunks)}")

    logger.info(f"Ingestion complete — {total_inserted} chunks in Supabase")
    return total_inserted


def main():
    parser = argparse.ArgumentParser(
        description="Ingest Colorix knowledge base into Supabase pgvector"
    )
    parser.add_argument(
        "--path", type=str, default=None,
        help="Path to .txt or .docx (overrides KNOWLEDGE_BASE_PATH in .env)",
    )
    args = parser.parse_args()

    from app.config import get_settings
    settings  = get_settings()
    docx_path = Path(args.path) if args.path else Path(settings.knowledge_base_path)

    print()
    print("═" * 62)
    print("  COLORIX GROUPE — Supabase Knowledge Base Ingestion")
    print("═" * 62)
    print(f"  Source  : {docx_path.resolve()}")
    print(f"  Supabase: {settings.supabase_url}")
    print(f"  Model   : Gemini text-embedding-004 (768-dim)")
    print(f"  Chunks  : size={settings.chunk_size}, overlap={settings.chunk_overlap}")
    print("═" * 62)
    print()

    start   = time.time()
    total   = run_ingestion(docx_path)
    elapsed = time.time() - start

    from app.database import document_count
    stored = document_count()

    print()
    print("═" * 62)
    print("INGESTION COMPLETE")
    print("═" * 62)
    print(f"  Chunks ingested : {total}")
    print(f"  Supabase total  : {stored}")
    print(f"  Time taken      : {elapsed:.1f}s")
    print("═" * 62)
    print()
    print("Start the server:")
    print("   uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload")
    print()


if __name__ == "__main__":
    main()